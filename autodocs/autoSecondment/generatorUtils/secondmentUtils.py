import os
import docx
import json
import openpyxl
import datetime

from docx import Document
from autoSecondment.generatorUtils.SecondmentData import SecondmentData

def getCountry (cell_country):
    return cell_country.split(" ")[0]
    

def getGalininkaVarda(full_name, linksniai_dic):
    names_split = full_name.strip().split(" ")
    if len(names_split) != 2:
        return False

    
    forename = names_split[1]
    surname = names_split[0]
    try:
        if forename[-1] == "s":
            forename = forename[:-2] + linksniai_dic[forename[-2:]]
        else: 
            forename = forename[:-1] + linksniai_dic[forename[-1:]]

        if surname[-1] == "s":
            surname = surname[:-2] + linksniai_dic[surname[-2:]]
        else: 
            surname =  surname[:-1] + linksniai_dic[surname[-1:]]
    except:
        return False
    return forename + " " + surname


def getXlsxData (file_name, linksniai_dic, INVALID_GA, INVALID_SPEC, COUNTRIES):
    records = []
    unfit_records = []
    START_ROW = 8 #TODO: input from cmd
    FINISH_ROW = START_ROW + int(input("--> Įveskite sveikąjį skaičių: ")) #TODO: automaticly stop after finding blank row
#    INVALID_GA = ["ADM","GTS"]
#    INVALID_SPEC = ["buh.","pavad.","sauga","techn.", "vad.", "vadov.", "vadyb.", "pat"]
#    COUNTRIES = ["Švedija", "Prancūzija", "Belgija", "Šveicarija", "Vokietija", "Latvija", "Ryga", "Lenkija", "Estija", "Nyderlandai", "Olandija", "Suomija"]
    wb = openpyxl.load_workbook(file_name)
    sheet = wb[wb.sheetnames[0]]
    
    for row in range(START_ROW, FINISH_ROW+1):
        this = []
        found_invalid = False
        
        nr = sheet[("A%s" % row)].value
        date_created = sheet[("B%s" % row)].value
        surname_name =  sheet[("C%s" % row)].value
        place = sheet[("D%s" % row)].value
        specialybe = sheet[("E%s" % row)].value
        country = sheet[("F%s" % row)].value
        date_from = sheet[("H%s" % row)].value
        date_to = sheet[("I%s" % row)].value

        if nr != None:

            if surname_name != None:
                surname_name = surname_name.strip()

            if type(nr) == str:
                nr = nr.strip()
            else:
                found_invalid = True
                
            if type(place) == str:
                place = place.strip()
            else:
                found_invalid = True

            if type(specialybe) == str:
                specialybe = specialybe.strip()
            else:
                found_invalid = True

            if type(country) == str:
                country = country.strip()
            else:
                found_invalid = True            
                
            galininkas_full_name = getGalininkaVarda(surname_name, linksniai_dic)
            if len(str(date_from)) > 10:
                date_from = str(date_from)[:10]

            if len(str(date_to)) > 10:
                date_to = str(date_to)[:10]

            if len(str(date_created)) > 10:
                date_created = str(date_created)[:10]


            if galininkas_full_name == False:
                #found_invalid = True
                galininkas_full_name = surname_name

            if country.strip().split(" ")[0] not in COUNTRIES:
                found_invalid = True
      
            this = [nr,date_created, surname_name, place, specialybe, country, date_from, date_to, galininkas_full_name]

            # probably dates are fcing this up. debug if all the field got correct type values
            if not found_invalid:
                for i in this:
                    #if i == "":
                    #    found_invalid = True
                    if i == None:
                        found_invalid = True
                if not found_invalid:
                    if specialybe in INVALID_SPEC:
                        found_invalid = True
                    elif place in INVALID_GA:
                        found_invalid = True
            
            if found_invalid:
                unfit_records.append(this)
            else:
                records.append(this)

    print('\n>>> Colected records from xlsx...\n')

    return records, unfit_records

def indexParagrahps(c):
    """ Find all placeholder detailed positions in a template document, and return them as dictionary, where key is placeholder key, and value is run number"""
    
    doc = Document(c.TEMPLATE_NAME)
    paragraphIdx = {}
    for pNum in range(len(doc.paragraphs)):
        for key in c.KEYS:
            if key in doc.paragraphs[pNum].text:
                paragraphIdx[key] = pNum
    return paragraphIdx

def indexRuns(c, paragraphIdx):
    """ Find all placeholder detailed positions in template paragraphs, and return them as dictionary, where key is placeholder key, and value is run number"""
    doc = Document(c.TEMPLATE_NAME)
    runIdx = {}
    for key in c.KEYS:
        for rNum in range(len(doc.paragraphs[paragraphIdx[key]].runs)):
            if key in doc.paragraphs[paragraphIdx[key]].runs[rNum].text:
                runIdx[key] = rNum
    return runIdx

def createSecondmentDocument (record, doc, paragraphIdx, runIdx, c, goodRecord):
    """[nr,date_created, surname_name, place, specialybe, country, date_from, date_to, galininkas_full_name]"""
    if goodRecord:
        docDir = c.DOCUMENTS_DIR
    else:
        docDir = c.FAILED_DOCUMENTS_DIR
    
    #doc = Document(c.TEMPLATE_NAME)
    s = SecondmentData(record)
    s.setConstants(c.GALININKO_LINKSNIAI, c.MONTHS, c.GA1, c.GA2, c.PROFESIJA, c.DIENPINIGIAI, c.KEYS)

#try:
    for key in c.KEYS:
        placeHolderRun = doc.paragraphs[paragraphIdx[key]].runs[runIdx[key]].text
        # print("key={}, s.translateKeyToValue(key)={}".format(key, s.translateKeyToValue(key)))
        doc.paragraphs[paragraphIdx[key]].runs[runIdx[key]].text = placeHolderRun.replace(key, s.translateKeyToValue(key))
#except Exception as e:
#    print('Something unrecognised for: %s\nError message: %s\n' % (record[2], repr(e)))
#    return
    currentDir = os.getcwd()
    
    if currentDir[-len(s.getCountry()):] != s.getCountry().upper():       
        try:
            os.chdir(c.FILE_DIR + '/' + docDir)
            os.chdir(s.getCountry().upper())
        except FileNotFoundError:
            os.chdir(c.FILE_DIR + '/' + docDir)
            os.mkdir(s.getCountry().upper())
            os.chdir(s.getCountry().upper())
    fileNr = ""
    fileExists = os.path.isfile(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
    if not fileExists:
        doc.save(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
    else:
        fileNr = 0
        while fileExists:
            fileNr += 1 
            fileExists = os.path.isfile(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
            if not fileExists:
                doc.save(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx')

def loadJsonData():
    with open('autoSecondment/generatorUtils/GALININKO_LINKSNIAI.json', 'r', encoding='utf8') as fp:
        GALININKO_LINKSNIAI = json.load(fp)
        
    with open('autoSecondment/generatorUtils/MONTHS.json', 'r', encoding='utf8') as fp:
        MONTHS = json.load(fp)

    with open('autoSecondment/generatorUtils/GA1.json', 'r', encoding='utf8') as fp:
        GA1 = json.load(fp)

    with open('autoSecondment/generatorUtils/GA2.json', 'r', encoding='utf8') as fp:
        GA2 = json.load(fp)

    with open('autoSecondment/generatorUtils/PROFESIJA.json', 'r', encoding='utf8') as fp:
        PROFESIJA = json.load(fp)

    with open('autoSecondment/generatorUtils/DIENPINIGIAI.json', 'r', encoding='utf8') as fp:
        DIENPINIGIAI = json.load(fp)
        
    with open('autoSecondment/generatorUtils/DocumentPlaceHolders.json', 'r', encoding='utf8') as fp:
        KEYS = json.load(fp)

    return GALININKO_LINKSNIAI, MONTHS, GA1, GA2, PROFESIJA, DIENPINIGIAI, KEYS

def mapCodeToCountry(key):
    countriesMap = {
    'ŠVED': 'Švedija',
    'PRAN': 'Prancūzija',
    'BELG': 'Belgija',
    'VOKE': 'Vokietija',
    'OLAND': 'Olandija',
    'Šveicarija': 'Šveicarija',
    'Latvija': 'Latvija',
    'Lenkija': 'Lenkija',
    'Estija': 'Estija',
    'Danija': 'Danija',
    'Slovakija': 'Slovakija',
    'Suomija': 'Suomija'
}

    try:
        return countriesMap[str(key).upper()]
    except KeyError:
        return key

def calculateLastWorkingDate(isoDate):
    if not type(isoDate) == str:
        return None
    currentDate = datetime.date.fromisoformat(isoDate)
    # currentDate = datetime.date(year=y, month=m, day=d)
    lastWorkingDate = ''
    foundDate = False
    minusDays = 1
    holidays = [
        "01-01",
        "02-16",
        "03-11",
        "05-01",
        "06-24",
        "07-06",
        "08-15",
        "11-01",
        "12-24",
        "12-25"
    ]
    while not foundDate:
        lastWorkingDate = currentDate - datetime.timedelta(days=minusDays)
        if lastWorkingDate.weekday() < 5 and lastWorkingDate.isoformat()[5:] not in holidays:
            foundDate = True
            continue
        else:
            minusDays += 1
    return lastWorkingDate.isoformat()